import time
from docx import Document as DocxDocument
from app.parsing.base import BaseParser
from app.parsing.schema import ParsedDocument, PageInfo, Block, ParsedTable, TableCell

class DocxParser(BaseParser):
    name = "docx_parser"
    version = "1.0"

    async def parse(self, document_id: str, content: bytes) -> ParsedDocument:
        import io
        t0 = time.perf_counter()
        doc = DocxDocument(io.BytesIO(content))
        blocks, tables, raw_text_parts = [], [], []
        order = 0

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            block_type = "heading" if para.style.name.startswith("Heading") else "paragraph"
            blocks.append(Block(block_id=f"{document_id}_b{order}", type=block_type, text=para.text, page=1, order=order))
            raw_text_parts.append(para.text)
            order += 1

        for t_idx, table in enumerate(doc.tables):
            cells = [
                TableCell(row=r_idx, col=c_idx, text=cell.text)
                for r_idx, row in enumerate(table.rows)
                for c_idx, cell in enumerate(row.cells)
            ]
            tables.append(ParsedTable(table_id=f"{document_id}_t{t_idx}", page=1, order=order, cells=cells))
            order += 1

        return ParsedDocument(
            document_id=document_id, source_format="docx", parser_name=self.name, parser_version=self.version,
            pages=[PageInfo(page_number=1)], blocks=blocks, tables=tables, images=[],
            raw_text="\n".join(raw_text_parts), reading_order_applied=True, ocr_used=False,
            processing_time_ms=(time.perf_counter() - t0) * 1000,
        )
